"""Read a clinical note out of whatever the hospital actually hands you: PDF, DOCX or text.

The gateway's job is to find identifiers in text, so everything upstream of it is one function:
``read(path) -> str``. The only interesting decisions are about *layout*, and they are not
cosmetic -- this project's own test set says PHI hides in headers, footers, signature blocks and
table cells, so an extractor that drops those is a leak before the pipeline ever runs:

* **Line and column structure is preserved.** ``detectors/structural.py`` classifies a line as a
  form field or table row by looking at runs of whitespace and colons. Reflowing a PDF into
  paragraphs destroys exactly that signal, and with it the stricter thresholds the structural
  pre-pass applies to high-PHI-density regions.
* **DOCX headers and footers are read explicitly.** ``python-docx`` exposes them only via
  ``section.header`` / ``section.footer``; iterating ``document.paragraphs`` silently skips them.
  A running header saying "Fitzsimmons, Clementine -- Page 1 of 3" is PHI on every page.
* **DOCX tables are read explicitly** for the same reason -- table cells are not paragraphs.

``ponytail:`` no OCR. A scanned PDF with no text layer yields nothing, and this raises rather
than returning "" so that a scan is never mistaken for a clean note. Add pytesseract behind the
same function if scanned intake forms ever become real input.
"""

from __future__ import annotations

from pathlib import Path

#: Suffixes handled without a converter. Anything else raises with the list, because guessing
#: at an unknown container and silently producing partial text is how PHI escapes.
SUFFIXES = (".txt", ".text", ".md", ".rtf", ".pdf", ".docx")


def _pdf(path: Path) -> str:
    """Page text, joined with a form feed so page boundaries survive into the pipeline.

    pdfplumber first: it keeps whitespace runs roughly where they were on the page, which is
    what the form-field and table heuristics downstream key off. pypdf is the fallback because
    it is the declared dependency and always present.
    """
    try:
        import pdfplumber
    except ImportError:
        pass
    else:
        with pdfplumber.open(str(path)) as pdf:
            pages = [(p.extract_text(layout=True) or "") for p in pdf.pages]
        if any(p.strip() for p in pages):
            return "\f".join(pages)

    from pypdf import PdfReader

    pages = [(p.extract_text() or "") for p in PdfReader(str(path)).pages]
    if not any(p.strip() for p in pages):
        raise ValueError(
            f"{path.name}: no text layer found. This looks like a scan; OCR it first "
            "(see the ponytail note in phi_gateway/ingest.py). Refusing to return empty text, "
            "because an empty note masks clean and tells you nothing."
        )
    return "\f".join(pages)


def _docx(path: Path) -> str:
    import docx  # python-docx

    d = docx.Document(str(path))
    out: list[str] = []

    for section in d.sections:
        for part in (section.header, section.footer):
            if part is None:
                continue
            for p in part.paragraphs:
                if p.text.strip():
                    out.append(p.text)

    for p in d.paragraphs:
        out.append(p.text)

    for table in d.tables:
        for row in table.rows:
            # Tab-joined, not space-joined: the structural pre-pass reads a wide gap as a
            # column break, and a single space would fuse "MRN" to its value.
            out.append("\t".join(c.text.strip() for c in row.cells))

    return "\n".join(out)


def read(path: str | Path) -> str:
    """Text of ``path``, dispatched on suffix. Raises on anything it cannot read honestly."""
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(p)
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return _pdf(p)
    if suffix == ".docx":
        return _docx(p)
    if suffix == ".doc":
        raise ValueError(
            f"{p.name}: legacy .doc is a binary OLE format, not a zip of XML. Convert to .docx "
            "or .txt first (`soffice --convert-to docx`)."
        )
    if suffix in SUFFIXES:
        return p.read_text(encoding="utf-8", errors="replace")
    # Unknown suffix: try text rather than refuse, but only if it decodes as text. A note
    # emailed as ".note" or ".dat" is common; a JPEG is not something to hand the tagger.
    raw = p.read_bytes()
    if b"\x00" in raw[:8192]:
        raise ValueError(
            f"{p.name}: binary content and unknown suffix {suffix!r}. Handled: "
            f"{', '.join(SUFFIXES)}"
        )
    return raw.decode("utf-8", errors="replace")


def _selfcheck() -> None:
    """One round trip per format, asserting the layout guarantees the docstring claims."""
    import tempfile

    with tempfile.TemporaryDirectory() as td:
        d = Path(td)

        t = d / "n.txt"
        t.write_text("MRN:  00-4471-882\nPatient: Wood, J\n", encoding="utf-8")
        assert "00-4471-882" in read(t)
        assert read(t).count("\n") == 2, "line structure must survive"

        u = d / "n.note"
        u.write_bytes(b"Seen by Halvorsen at 0620.\n")
        assert "Halvorsen" in read(u), "unknown-but-textual suffix should still read"

        b = d / "scan.bin"
        b.write_bytes(b"\x00\x01\x02" * 10)
        try:
            read(b)
            raise AssertionError("binary blob must raise, not decode to mojibake")
        except ValueError:
            pass

        try:
            import docx
        except ImportError:
            print("ingest selfcheck OK (docx skipped -- not installed)")
            return
        doc = docx.Document()
        doc.sections[0].header.paragraphs[0].text = "Fitzsimmons, Clementine -- Page 1"
        doc.add_paragraph("Echo shows preserved EF.")
        tbl = doc.add_table(rows=1, cols=2)
        tbl.rows[0].cells[0].text = "MRN"
        tbl.rows[0].cells[1].text = "8830192"
        w = d / "n.docx"
        doc.save(str(w))
        got = read(w)
        assert "Fitzsimmons" in got, "header PHI dropped -- the exact failure this guards"
        assert "8830192" in got, "table cell dropped"
        assert "MRN\t8830192" in got, "table columns must stay separated"
        print("ingest selfcheck OK")


if __name__ == "__main__":
    _selfcheck()
